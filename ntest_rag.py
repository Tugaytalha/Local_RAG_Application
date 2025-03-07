import time
from docx import Document
import umap
import numpy as np
import matplotlib.pyplot as plt
from run_utils import populate_database, evaluate_response, QueryData, get_embedding_function, get_all_chunk_embeddings
import warnings
warnings.filterwarnings("ignore")

RESULTS_PATH = "results/"

FROM_XLSX = True
XLSX_PATH = "queries.xlsx"


QUERIES_CUSTOMER = {
    "müşteriler bankamız ATMleri harici hangi ATMleri kullanabilir?":
        ["PTT", ["data\Çağrı merkezi chatbot için bilgiler v2.docx:None:44"]],
    "Hareketsiz hesap nedir?":
        [
            "Müşterimizin 1 yılı aşkın süre zarfında hesabına ilişkin para çıkışı yapmaması durumunda hesap hareketsiz konuma alınmaktadır. Ek9",
            ["data\Çağrı merkezi chatbot için bilgiler v2.docx:None:49"]],
    "Debit Kart kesin olarak kapatılmak istendiğinde Çağrı Merkezi personeli nasıl bir yol izlemelidir? > Hangi durumlarda Debit Kart geçici olarak kapatılır?":
        ["İstenen kartı kesin olarak kapatsa da derhal yenisinin başvurusunu almalıdır. > Ek8",
         ["data\Çağrı merkezi chatbot için bilgiler v2.docx:None:31",
          "data\Çağrı merkezi chatbot için bilgiler v2.docx:None:9"]],

    # "müşterinin hangi durumlarda para transferi işlemini Çağrı Merkezinden yapması mümkündür?": [
    #     "Son 20 işlem veya kayıtlı işlemleri Çağrı Merkezinden yapabilir. Güvenlik gereği başka para transferi işlemlerini yapamaz.",
    #     [""]],

    "kayıp çalışntı durumunda kapatılan karta bağlı HGS talimatı yeni verilen karta otomatik devrolur mu?": [
        "Evet",
        ["data\Çağrı merkezi chatbot için bilgiler v2.docx:None:11"]],
    "ATM'lerde yapılan işlemlerde hangi koşullarda müşteriden komisyon alınır?": [
        "Bankamız ATM'lerinde bankamız kartı ile yapılan hiçbir işlemde komisyon alınmaz, başka banka ATMlerinden işlem yapılması halinde komisyon alınır. Ek7",
        ["data\Çağrı merkezi chatbot için bilgiler v2.docx:None:45", "data\Çağrı merkezi chatbot için bilgiler v2.docx:None:46"]],
    "Kredi kartım suya düşse ne olur?": ["Buna cevap veremiyorum.", []],
}

QUERIES_FINANCIAL = {
    "30 Eylül 2024 itibarıyla Bankanın toplam varlıkları ne kadardır?": [
        "VARLIKLAR TOPLAMI 152.836.357 144.170.432 297.006.789 123.217.699 108.951.796 232.169.495",
        [] #"1.pdf:None:101","1.pdf:None:102"
    ],
    "30 Eylül 2024 itibarıyla Bankanın nakdi kredileri ne kadardır?": [
        "2.1 Krediler  (6) 77.433.485 51.342.362 128.775.847 68.988.788 35.631.066 104.619.854",
        [] #"1.pdf:None:42"
    ],
    "30 Eylül 2024 itibarıyla Bankanın itfa edilmiş maliyeti ile ölçülen finansal varlıkları ne kadardır?": [
        "2.3 İtfa Edilmiş Maliyeti ile Ölçülen Finansal Varlıklar (4) 12.715.079 16.363.818 29.078.897 11.533.660 17.318.883 28.852.543",
        [] #"1.pdf:None:43","1.pdf:None:44"
    ],
    # "30 Eylül 2024 itibarıyla Bankanın maddi duran varlıkları (net) ne kadardır?": [
    #     "V. MADDİ DURAN VARLIKLAR (Net)  (10) 5.094.234 79.469 5.173.703 4.448.035 66.241 4.514.276",
    #     [] #"1.pdf:None:51"
    # ],
    # "30 Eylül 2024 itibarıyla Bankanın toplanan fonlar ne kadardır?": [
    #     "YÜKÜMLÜLÜKLER I. TOPLANAN FONLAR (1) 103.207.064 93.228.495 196.435.559 81.304.541 80.501.315 161.805.856",
    #     [] #"1.pdf:None:55"
    # ],
    # "30 Eylül 2024 itibarıyla Bankanın alınan krediler ne kadardır?": [
    #     "II. ALINAN KREDİLER (2) 14.156.677 40.012.913 54.169.590 2.178.308 28.357.631 30.535.939",
    #     [] #"1.pdf:None:56"
    # ],
    #  "30 Eylül 2024 itibarıyla Bankanın özkaynakları ne kadardır?": [
    #     "XIV. ÖZKAYNAKLAR (10) 15.923.036 135.682 16.058.718 13.326.608 60.811 13.387.419",
    #     [] #"1.pdf:None:62","1.pdf:None:63"
    # ],
    #  "2024 yılı 3. çeyrekte kar payı gelirleri ne kadar gerçekleşmiştir?": [
    #     "I. KÂR PAYI GELİRLERİ (1) 28.879.308 11.541.042 10.614.731 4.608.499",
    #     [] #"2.pdf:None:1"
    # ],
    # "2024 yılı 3. çeyrekte kar payı giderleri ne kadar gerçekleşmiştir?": [
    #     "II. KAR PAYI GİDERLERİ (-) (2) 23.419.002 6.652.405 9.869.893 2.733.651",
    #     [] #"2.pdf:None:6"
    # ],
    # "2024 yılı 3. çeyrekte personel giderleri ne kadar gerçekleşmiştir?": [
    #     "XI. PERSONEL GIDERLERI (-) (6) 3.463.260 1.836.508 1.011.553 633.109",
    #     [] #"2.pdf:None:44"
    # ],
    # "Bankanın 2024 yılı üçüncü çeyrekteki dönem net karı ne kadardır?": [
    #     "XXV. DÖNEM NET KARI/ZARARI (XIV+XXIV) (12) 2.664.097 2.460.503 793.617 967.310",
    #     [] #"2.pdf:None:49"
    # ]

}

QUERIES = QUERIES_CUSTOMER

EMBEDDING_MODELS = [
    #####"emrecan/convbert-base-turkish-mc4-cased-allnli_tr",
    ###"emrecan/bert-base-turkish-cased-mean-nli-stsb-tr",
    ###"atasoglu/roberta-small-turkish-clean-uncased-nli-stsb-tr",
    ###"atasoglu/distilbert-base-turkish-cased-nli-stsb-tr",
    ## "atasoglu/xlm-roberta-base-nli-stsb-tr",
    ## "atasoglu/mbert-base-cased-nli-stsb-tr",
    # "jinaai/jina-embeddings-v3",
    "Omerhan/intfloat-fine-tuned-14376-v4",
    "intfloat/multilingual-e5-large-instruct",
    ###"atasoglu/turkish-base-bert-uncased-mean-nli-stsb-tr",
]


def visualize_queries(queries, query_embeddings, all_chunk_embeddings, retrieved_embeddings_list, expected_embeddings_list, model_name):
    """
    Creates subplots for each query, visualizing its relationship with all chunks, retrieved chunks, and expected chunks.
    """
    num_queries = len(queries)
    cols = 2  # Number of columns for subplots
    rows = (num_queries + 1) // cols  # Compute required rows dynamically

    # Reduce dimensions using UMAP
    reducer = umap.UMAP(n_neighbors=15, min_dist=0.1, metric='cosine')
    all_embeddings = np.vstack((query_embeddings, all_chunk_embeddings))
    reduced_embeddings = reducer.fit_transform(all_embeddings)

    # Extract 2D positions
    query_positions = reduced_embeddings[:len(query_embeddings)]
    all_chunk_positions = reduced_embeddings[len(query_embeddings):]

    fig, axes = plt.subplots(rows, cols, figsize=(12, 6 * rows))  # Dynamic plot size
    axes = axes.flatten()

    for i, (query, query_embedding) in enumerate(zip(queries, query_positions)):
        ax = axes[i]

        # Plot all chunks in grey
        ax.scatter(all_chunk_positions[:, 0], all_chunk_positions[:, 1], c='grey', label='All Chunks', alpha=0.3)

        # Plot query in red
        ax.scatter(query_embedding[0], query_embedding[1], c='red', label='Query', s=100, edgecolor='black')

        # Plot retrieved chunks in yellow
        retrived_to_pilot = reducer.transform(retrieved_embeddings_list[i])
        ax.scatter(retrived_to_pilot[:, 0], retrived_to_pilot[:, 1], c='yellow', label='Retrieved Chunks', s=70,
                   edgecolor='black')

        # Plot expected chunks in green
        if len(expected_embeddings_list[i]) != 0:
            expected_to_pilot = reducer.transform(expected_embeddings_list[i])
            ax.scatter(expected_to_pilot[:, 0], expected_to_pilot[:, 1], c='green', label='Expected Chunks', s=70,
                       edgecolor='black')

        # Decide expected chunks retrieved or not
        expectation_list = [exp in retrieved_embeddings_list[i] for exp in expected_embeddings_list[i]]
        ax.set_title(f"Query {i + 1}: {query[:40]}{'...' if len(query) > 40 else ''}, R1:{any(expectation_list)}, RA:{all(expectation_list)}")  # Show first 40 chars of the query
        ax.legend(loc='upper right')
        ax.grid(True)

    plt.tight_layout()
    plt.savefig(RESULTS_PATH + (f"query_chunks_visualization_{model_name}.png").replace(" ", "_").replace("/", "_"))
    plt.close()


def try_rag_with_embeddings(embedding_model_name):
    """
    Tests the RAG application with a specific embedding model.
    """
    start_time = time.time()

    # Create a Word document to store the test results
    document = Document()
    document.add_heading("RAG Application Test Report", 0)

    print(f"Testing with embedding: {embedding_model_name}")
    document.add_heading(f"Embedding: {embedding_model_name}", level=1)

    # Populate the database with the specified embedding model
    populate_database(reset=True, model_name=embedding_model_name, model_type="sentence_transformer")
    database_end_time = time.time()
    print(f"Database populated in {database_end_time - start_time:.2f} seconds")

    # Initialize embedding function with appropriate settings
    embedding = get_embedding_function(
        model_name_or_path=embedding_model_name,
        model_type="sentence_transformer"
    )

    # Get all chunk embeddings from the database
    all_chunk_data = get_all_chunk_embeddings()
    all_chunk_embeddings = np.array(all_chunk_data["embeddings"])

    query_embeddings = []
    retrieved_embeddings_list = []
    expected_embeddings_list = []
    all_sources = []

    # Test with each query
    for query, expected_response_chunk in QUERIES.items():
        query_embedding = np.array(embedding.embed_query(query)).reshape(1, -1)
        query_embeddings.append(query_embedding)

        try:
            response, retrieved_chunks = QueryData.query_rag(query, embedding)
        except Exception as e:
            print(f"Error during query processing: {e}")
            response = "Error during query processing"
            retrieved_chunks = []

        retrieved_embeddings = np.array([embedding.embed_query(chunk['content']) for chunk in retrieved_chunks])
        retrieved_embeddings_list.append(retrieved_embeddings)

        expected_embeddings = []
        for source in expected_response_chunk[1]:
            for chunk in all_chunk_data["ids"]:
                if chunk == source:
                    expected_embeddings.append(all_chunk_data["embeddings"][all_chunk_data["ids"].index(chunk)])
        expected_embeddings = np.array(expected_embeddings)
        expected_embeddings_list.append(expected_embeddings)

        # Take sources
        sources = [chunk['source'] for chunk in retrieved_chunks]

        # Add sources to the list
        all_sources.append(sources)

        evaluation_result = evaluate_response(response, expected_response_chunk[0])

        # Add results to the document
        document.add_paragraph(f"Query: {query}")
        document.add_paragraph(f"Expected Response: {expected_response_chunk[0]}")
        document.add_paragraph(f"Actual Response: {response}")
        document.add_paragraph(f"Evaluation: {evaluation_result}")
        document.add_paragraph(f"Sources: {sources}")
        document.add_paragraph("Retrieved Chunks:")
        for chunk in retrieved_chunks:
            document.add_paragraph(
                f" - Source: {chunk['source']}, Content: {chunk['content']}, Score: {chunk['score']}")
        document.add_paragraph("---")

    visualize_queries(QUERIES.keys(), np.vstack(query_embeddings), all_chunk_embeddings, retrieved_embeddings_list,
                      expected_embeddings_list, embedding_model_name)

    # Add sources to the end of the document
    document.add_heading("Sources", level=1)
    for sources in all_sources:
        document.add_paragraph(f"Sources: {sources}")

    end_time = time.time()
    print(f"Test completed in {end_time - start_time:.2f} seconds")
    document.add_paragraph(f"Test completed in {end_time - start_time:.2f} seconds")

    document.save(RESULTS_PATH + f"rag_test_report_{embedding_model_name}.docx".replace("/", "_"))
    print(f"RAG test report generated: rag_test_report_{embedding_model_name}.docx")


def main():
    """
    Runs the RAG tests for all embedding models.
    """
    if FROM_XLSX:
        global QUERIES
        import pandas as pd
        # Read queries from the Excel file
        queries_df = pd.read_excel(XLSX_PATH)
        # Use question col as key and answer col + empty list as value
        QUERIES = dict(zip(queries_df["question"], queries_df["answer"].apply(lambda x: [x, []])))
        print(f"Queries read from the Excel file: {QUERIES}")

    for embedding_model_name in EMBEDDING_MODELS:
        try_rag_with_embeddings(embedding_model_name)


if __name__ == "__main__":
    main()