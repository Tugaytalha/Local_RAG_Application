from docx import Document
from run_utils import populate_database, evaluate_response
from get_embedding_function import get_embedding_function
from populate_database import clear_database


TEST_QUERIES = {
    "müşteriler bankamız ATMleri harici hangi ATMleri kullanabilir?": "PTT",
    "Hareketsiz hesap nedir?": "Müşterimizin 1 yılı aşkın süre zarfında hesabına ilişkin para çıkışı yapmaması durumunda hesap hareketsiz konuma alınmaktadır. Ek9",
    "Debit Kart kesin olarak kapatılmak istendiğinde Çağrı Merkezi personeli nasıl bir yol izlemelidir? > Hangi durumlarda Debit Kart geçici olarak kapatılır?": "İstenen kartı kesin olarak kapatsa da derhal yenisinin başvurusunu almalıdır. > Ek8",
    "müşterinin hangi durumlarda para transferi işlemini Çağrı Merkezinden yapması mümkündür?": "Son 20 işlem veya kayıtlı işlemleri Çağrı Merkezinden yapabilir. Güvenlik gereği başka para transferi işlemlerini yapamaz.",
    "kayıp çalışntı durumunda kapatılan karta bağlı HGS talimatı yeni verilen karta otomatik devrolur mu?": "Evet",
    "ATM'lerde yapılan işlemlerde hangi koşullarda müşteriden komisyon alınır?": "Bankamız ATM'lerinde bankamız kartı ile yapılan hiçbir işlemde komisyon alınmaz, başka banka ATMlerinden işlem yapılması halinde komisyon alınır. Ek7",
    "Kredi kartım suya düşse ne olur?": "Buna cevap veremiyorum.",
}

EMBEDDING_MODELS = [
    # "emrecan/convbert-base-turkish-mc4-cased-allnli_tr",
    # "emrecan/bert-base-turkish-cased-mean-nli-stsb-tr",
    # "atasoglu/roberta-small-turkish-clean-uncased-nli-stsb-tr",
    # "atasoglu/distilbert-base-turkish-cased-nli-stsb-tr",
    # "atasoglu/xlm-roberta-base-nli-stsb-tr",
    # "atasoglu/mbert-base-cased-nli-stsb-tr",
    # "Omerhan/intfloat-fine-tuned-14376-v4",
    "atasoglu/turkish-base-bert-uncased-mean-nli-stsb-tr",
    # "jinaai/jina-embeddings-v3",
]

def test_rag_with_embeddings(embedding_model_name):
    """
    Tests the RAG application with a specific embedding model.
    """
    # Create a new DOCX document for the report
    document = Document()
    document.add_heading("RAG Application Test Report", 0)

    print(f"Testing with embedding: {embedding_model_name}")
    document.add_heading(f"Embedding: {embedding_model_name}", level=1)

    # Populate the database
    populate_database(reset=True, model_name=embedding_model_name, model_type="sentence_transformer")

    # Get the embedding function
    embedding = get_embedding_function(embedding_model_name, "sentence_transformer")

    all_sources = []
    # Test with each query
    for query, expected_response in TEST_QUERIES.items():
        try:
            from run_utils import query_rag
            response, retrieved_chunks = query_rag(query, embedding)
        except Exception as e:
            print(f"Error during query processing: {e}")
            response = "Error during query processing"
            retrieved_chunks = []

        evaluation_result = evaluate_response(response, expected_response)

        # Take sources
        sources = [chunk['source'] for chunk in retrieved_chunks]

        # Add sources to the list
        all_sources.append(sources)

        # Add results to the document
        document.add_paragraph(f"Query: {query}")
        document.add_paragraph(f"Expected Response: {expected_response}")
        document.add_paragraph(f"Actual Response: {response}")
        document.add_paragraph(f"Evaluation: {evaluation_result}")
        document.add_paragraph(f"Sources: {sources}")

        # Add retrieved chunks to the document
        document.add_paragraph("Retrieved Chunks:")
        for chunk in retrieved_chunks:
            document.add_paragraph(f" - Source: {chunk['source']}, Content: {chunk['content']}, Score: {chunk['score']}")

        document.add_paragraph("---")

    # Add sources to the end of the document
    document.add_heading("Sources", level=1)
    for sources in all_sources:
        document.add_paragraph(f"Sources: {sources}")

    # Save the document
    document.save((f"rag_test_report_{embedding_model_name}.docx").replace("/", "_"))
    print(f"RAG test report generated: rag_test_report_{embedding_model_name}.docx")


def main():
    """
    Runs the RAG tests for all embedding models.
    """
    for embedding_model_name in EMBEDDING_MODELS:
        test_rag_with_embeddings(embedding_model_name)

if __name__ == "__main__":
    main()